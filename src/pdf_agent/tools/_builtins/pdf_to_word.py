"""PDF to Word tool — convert PDF to DOCX using LibreOffice."""
from __future__ import annotations

import shutil
from pathlib import Path

import pikepdf

from pdf_agent.core import ErrorCode, ToolError
from pdf_agent.schemas.tool import ToolInputSpec, ToolManifest, ToolOutputSpec
from pdf_agent.tools._builtins.pdf_to_text import _extract_page_text
from pdf_agent.tools.base import BaseTool, ProgressReporter, ToolResult
from pdf_agent.tools.filenames import localized_output_name
from pdf_agent.tools.libreoffice import run_libreoffice_conversion_to_output


class PdfToWordTool(BaseTool):
    def manifest(self) -> ToolManifest:
        return ToolManifest(
            name="pdf_to_word",
            label="PDF 转 Word",
            category="convert",
            description="将 PDF 转换为可编辑的 Word 文档（.docx）",
            inputs=ToolInputSpec(min=1, max=1),
            outputs=ToolOutputSpec(type="docx"),
            params=[],
            engine="libreoffice+python-docx",
            async_hint=True,
        )

    def validate(self, params: dict) -> dict:
        return {}

    def run(
        self,
        inputs: list[Path],
        params: dict,
        workdir: Path,
        reporter: ProgressReporter | None = None,
    ) -> ToolResult:
        output_path = workdir / localized_output_name(inputs[0], "转Word", ext=".docx")
        lo_bin = shutil.which("libreoffice") or shutil.which("soffice")
        engine = "python-docx"
        fallback_used = not bool(lo_bin)
        fallback_reason: str | None = "libreoffice not found" if not lo_bin else None
        if lo_bin:
            if reporter:
                reporter(10, "Starting LibreOffice conversion...")
            success, failure_reason = run_libreoffice_conversion_to_output(
                lo_bin,
                convert_to="docx",
                input_path=inputs[0],
                output_path=output_path,
                outdir=workdir,
                profile_dir=workdir / ".libreoffice-profile",
            )
            if success:
                engine = "libreoffice"
            else:
                fallback_used = True
                fallback_reason = failure_reason

        if engine != "libreoffice":
            self._fallback_convert(inputs[0], output_path)

        if reporter:
            reporter(100, "Conversion complete")

        meta = {
            "original": inputs[0].name,
            "output": output_path.name,
            "engine": engine,
            "fallback_used": fallback_used,
        }
        if fallback_reason:
            meta["fallback_reason"] = fallback_reason
        log = f"Converted {inputs[0].name} to {output_path.name} via {engine}"
        if fallback_reason:
            log += f" (fallback reason: {fallback_reason})"
        return ToolResult(
            output_files=[output_path],
            meta=meta,
            log=log,
        )

    @staticmethod
    def _fallback_convert(input_pdf: Path, output_path: Path) -> None:
        try:
            from docx import Document
        except ImportError as exc:
            raise ToolError(ErrorCode.ENGINE_NOT_INSTALLED, "python-docx is not installed") from exc

        document = Document()
        document.add_heading(input_pdf.stem, level=1)
        with pikepdf.open(input_pdf) as pdf:
            for page_index, page in enumerate(pdf.pages, start=1):
                document.add_heading(f"Page {page_index}", level=2)
                text = _extract_page_text(page).strip()
                if text:
                    for line in text.splitlines():
                        if line.strip():
                            document.add_paragraph(line.strip())
                else:
                    document.add_paragraph("[No extractable text]")
        document.save(output_path)
